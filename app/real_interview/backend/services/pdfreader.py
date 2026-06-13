import io
import os
import re
from datetime import datetime, timezone
from typing import Any, BinaryIO, Dict, List, Optional, Tuple, Union

import gridfs
from bson import ObjectId
from bson.errors import InvalidId
from dotenv import load_dotenv
from pymongo import MongoClient
from pymongo.collection import Collection
from pymongo.database import Database
from pymongo.errors import CollectionInvalid, PyMongoError
from pypdf import PdfReader
from docx import Document

from app.real_interview import logger
from app.real_interview.backend.agents.resume_parse_agent import ResumeParseAgent
from app.real_interview.backend.utils.mongodb import get_shared_mongodb_client

load_dotenv()


def _get_db_name() -> str:
    logger.info("inside _get_db_name")
    name = os.getenv("MONGODB_DB_NAME", "").strip()
    if not name:
        raise ValueError("MONGODB_DB_NAME is not set in the environment or .env file")
    return name


def _get_resume_collection_name() -> str:
    logger.info("inside _get_resume_collection_name")
    name = os.getenv("MONGODB_COLLECTION_RESUMES", "").strip()
    if not name:
        raise ValueError(
            "MONGODB_COLLECTION_RESUMES is not set in the environment or .env file"
        )
    return name


def _get_gridfs_bucket() -> str:
    logger.info("inside _get_gridfs_bucket")
    bucket = os.getenv("MONGODB_GRIDFS_RESUME_BUCKET", "").strip()
    return bucket or "resume_pdf_fs"


RESUME_EXTENSIONS = frozenset({".pdf", ".doc", ".docx"})
RESUME_CONTENT_TYPES = {
    "pdf": "application/pdf",
    "doc": "application/msword",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
DEFAULT_RESUME_FILENAMES = {
    "pdf": "resume.pdf",
    "doc": "resume.doc",
    "docx": "resume.docx",
}


def _file_extension(filename: Optional[str]) -> str:
    if not filename:
        return ""
    return os.path.splitext(filename.lower().strip())[1]


def detect_resume_format(*, data: bytes, filename: Optional[str] = None) -> str:
    """Return pdf, doc, or docx based on magic bytes and filename."""
    if data.startswith(b"%PDF"):
        return "pdf"
    if len(data) >= 4 and data[:4] == b"PK\x03\x04":
        return "docx"
    if len(data) >= 8 and data[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
        return "doc"

    ext = _file_extension(filename)
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext == ".doc":
        return "doc"
    raise ValueError("only PDF, DOC, and DOCX documents are allowed")


def _as_user_object_id(userid: Union[str, ObjectId]) -> ObjectId:
    logger.info("inside _as_user_object_id")
    if isinstance(userid, ObjectId):
        return userid
    if not isinstance(userid, str) or not userid.strip():
        raise ValueError("userid must be a non-empty string or ObjectId")
    try:
        return ObjectId(userid.strip())
    except InvalidId as exc:
        raise ValueError("userid is not a valid ObjectId") from exc


def _default_parsed_data() -> Dict[str, Any]:
    logger.info("inside _default_parsed_data")
    return {
        "name": "",
        "email": "",
        "phone": "",
        "education": [],
        "experience": [],
        "skills": [],
        "certifications": [],
    }


def _heuristic_parse(raw_text: str) -> Dict[str, Any]:
    logger.info("inside _heuristic_parse")
    data = _default_parsed_data()
    if not raw_text or not raw_text.strip():
        return data

    text = raw_text.strip()
    email_match = re.search(
        r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
        text,
    )
    if email_match:
        data["email"] = email_match.group(0)

    phone_match = re.search(
        r"(?:\+?\d{1,3}[-.\s]?)?(?:\(?\d{2,4}\)?[-.\s]?)?\d{2,4}[-.\s]?\d{2,4}[-.\s]?\d{2,6}",
        text,
    )
    if phone_match:
        data["phone"] = phone_match.group(0).strip()

    first_line = text.splitlines()[0].strip() if text.splitlines() else ""
    if first_line and len(first_line) <= 80 and "@" not in first_line:
        data["name"] = first_line

    return data


def _overlay_scalar_fields(target: Dict[str, Any], overlay: Dict[str, Any]) -> None:
    logger.info("inside _overlay_scalar_fields")
    for key in ("name", "email", "phone"):
        if not (target.get(key) or "").strip() and (overlay.get(key) or "").strip():
            target[key] = overlay[key]


class resume_reader:
    """
    Read resume documents (PDF, DOC, DOCX), validate uploads, persist binaries in GridFS,
    and store metadata plus extracted text in the resume collection from the environment.
    """

    def __init__(
        self,
        client: Optional[MongoClient] = None,
        parse_agent: Optional[ResumeParseAgent] = None,
    ) -> None:
        logger.info("inside __init__")
        self._owns_client = client is not None
        self._client: MongoClient = client or get_shared_mongodb_client()
        self._db: Database = self._client[_get_db_name()]
        self._collection_name = _get_resume_collection_name()
        self._gridfs_bucket = _get_gridfs_bucket()
        self._fs = gridfs.GridFS(self._db, collection=self._gridfs_bucket)
        self._parse_agent = parse_agent if parse_agent is not None else ResumeParseAgent()

    def close(self) -> None:
        logger.info("inside close")
        if self._owns_client and self._client is not None:
            self._client.close()

    def _collection(self) -> Collection:
        logger.info("inside _collection")
        return self._db[self._collection_name]

    def ensure_resume_collection(self) -> None:
        logger.info("inside ensure_resume_collection")
        names = set(self._db.list_collection_names())
        if self._collection_name not in names:
            try:
                self._db.create_collection(self._collection_name)
                logger.info(
                    "[resume_reader] created collection %r",
                    self._collection_name,
                )
            except CollectionInvalid:
                logger.warning(
                    "[resume_reader] collection %r already exists (race)",
                    self._collection_name,
                )
        else:
            logger.info(
                "[resume_reader] collection %r already present",
                self._collection_name,
            )

        coll = self._collection()
        coll.create_index([("userid", 1), ("uploaded_ts", -1)], name="userid_uploaded_ts")
        logger.info("[resume_reader] indexes ensured on %r", self._collection_name)

    @staticmethod
    def validate_resume_file(
        *,
        data: bytes,
        filename: Optional[str] = None,
        content_type: Optional[str] = None,
    ) -> str:
        """Validate upload and return detected format: pdf, doc, or docx."""
        logger.info("inside validate_resume_file")
        _ = content_type  # reserved for callers (e.g. Flask)
        if not data:
            raise ValueError("file size must be greater than zero")
        ext = _file_extension(filename)
        if ext and ext not in RESUME_EXTENSIONS:
            raise ValueError("only PDF, DOC, and DOCX documents are allowed")
        return detect_resume_format(data=data, filename=filename)

    @staticmethod
    def validate_pdf_file(
        *,
        data: bytes,
        filename: Optional[str] = None,
        content_type: Optional[str] = None,
    ) -> None:
        """Backward-compatible PDF-only validation."""
        fmt = resume_reader.validate_resume_file(
            data=data, filename=filename, content_type=content_type
        )
        if fmt != "pdf":
            raise ValueError("only PDF documents are allowed")

    @staticmethod
    def extract_raw_text_from_pdf_bytes(data: bytes) -> str:
        logger.info("inside extract_raw_text_from_pdf_bytes")
        try:
            reader = PdfReader(io.BytesIO(data))
            parts: List[str] = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
            return "\n".join(parts).strip()
        except Exception as exc:
            logger.exception("[resume_reader] failed to read PDF")
            raise ValueError("could not read PDF content") from exc

    @staticmethod
    def extract_raw_text_from_docx_bytes(data: bytes) -> str:
        logger.info("inside extract_raw_text_from_docx_bytes")
        try:
            doc = Document(io.BytesIO(data))
            parts: List[str] = []
            for para in doc.paragraphs:
                text = (para.text or "").strip()
                if text:
                    parts.append(text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = (cell.text or "").strip()
                        if text:
                            parts.append(text)
            return "\n".join(parts).strip()
        except Exception as exc:
            logger.exception("[resume_reader] failed to read DOCX")
            raise ValueError("could not read DOCX content") from exc

    @staticmethod
    def extract_raw_text_from_doc_bytes(data: bytes) -> str:
        logger.info("inside extract_raw_text_from_doc_bytes")
        try:
            from legacy_doc import extract_text as extract_legacy_doc_text
        except ImportError as exc:
            raise ValueError(
                "DOC upload requires Python 3.11 or newer. Use PDF or DOCX instead."
            ) from exc
        try:
            result = extract_legacy_doc_text(data)
            text = (result.text or "").strip()
            if not text:
                raise ValueError("document contained no readable text")
            return text
        except ValueError:
            raise
        except Exception as exc:
            logger.exception("[resume_reader] failed to read DOC")
            raise ValueError("could not read DOC content") from exc

    @staticmethod
    def extract_raw_text(data: bytes, fmt: str) -> str:
        if fmt == "pdf":
            return resume_reader.extract_raw_text_from_pdf_bytes(data)
        if fmt == "docx":
            return resume_reader.extract_raw_text_from_docx_bytes(data)
        if fmt == "doc":
            return resume_reader.extract_raw_text_from_doc_bytes(data)
        raise ValueError("unsupported resume format")

    def save_resume(
        self,
        userid: Union[str, ObjectId],
        data: bytes,
        original_filename: str,
        *,
        content_type: Optional[str] = None,
    ) -> Dict[str, Any]:
        logger.info("inside save_resume")
        fmt = self.validate_resume_file(
            data=data, filename=original_filename, content_type=content_type
        )
        user_oid = _as_user_object_id(userid)

        self.ensure_resume_collection()
        raw_text = self.extract_raw_text(data, fmt)
        parsed_data = self._parse_agent.populate_parsed_data(raw_text)
        _overlay_scalar_fields(parsed_data, _heuristic_parse(raw_text))

        safe_name = os.path.basename(original_filename) or DEFAULT_RESUME_FILENAMES[fmt]
        stored_content_type = content_type or RESUME_CONTENT_TYPES[fmt]
        file_id = self._fs.put(
            data,
            filename=safe_name,
            content_type=stored_content_type,
            metadata={"userid": user_oid, "uploaded_ts": datetime.now(timezone.utc)},
        )

        resume_file_ref: Dict[str, str] = {
            "storage": "gridfs",
            "bucket": self._gridfs_bucket,
            "file_id": str(file_id),
            "filename": safe_name,
        }

        uploaded_ts = datetime.now(timezone.utc)
        doc = {
            "userid": user_oid,
            "resume_file": resume_file_ref,
            "raw_text": raw_text,
            "parsed_data": parsed_data,
            "uploaded_ts": uploaded_ts,
        }

        try:
            result = self._collection().insert_one(doc)
        except PyMongoError:
            logger.exception("[resume_reader] insert failed; removing orphaned GridFS file")
            try:
                self._fs.delete(file_id)
            except Exception:
                logger.exception("[resume_reader] failed to delete GridFS file after insert error")
            raise

        out = {
            "_id": str(result.inserted_id),
            "userid": str(user_oid),
            "resume_file": resume_file_ref,
            "raw_text": raw_text,
            "parsed_data": parsed_data,
            "uploaded_ts": uploaded_ts,
        }
        logger.info("[resume_reader] stored resume for user %s", user_oid)
        return out

    def read_resume_stream(
        self,
        stream: BinaryIO,
        *,
        filename: Optional[str] = None,
        content_type: Optional[str] = None,
    ) -> bytes:
        logger.info("inside read_resume_stream")
        data = stream.read()
        self.validate_resume_file(data=data, filename=filename, content_type=content_type)
        return data

    def read_pdf_stream(
        self,
        stream: BinaryIO,
        *,
        filename: Optional[str] = None,
        content_type: Optional[str] = None,
    ) -> bytes:
        """Backward-compatible alias."""
        return self.read_resume_stream(
            stream, filename=filename, content_type=content_type
        )

    def list_resumes_for_user(self, userid: Union[str, ObjectId]) -> List[Dict[str, Any]]:
        """Return resume summaries for a user, newest upload first."""
        logger.info("inside list_resumes_for_user")
        user_oid = _as_user_object_id(userid)
        self.ensure_resume_collection()
        cursor = (
            self._collection()
            .find(
                {"userid": user_oid},
                {"uploaded_ts": 1, "parsed_data.name": 1, "resume_file.filename": 1},
            )
            .sort("uploaded_ts", -1)
        )
        items: List[Dict[str, Any]] = []
        for doc in cursor:
            ts = doc.get("uploaded_ts")
            parsed = doc.get("parsed_data") or {}
            resume_file = doc.get("resume_file") or {}
            filename = (resume_file.get("filename") or "").strip()
            name = (parsed.get("name") or "").strip() if isinstance(parsed, dict) else ""
            label = filename or name or "Resume"
            uploaded_ts = ts.isoformat() if hasattr(ts, "isoformat") else str(ts)
            items.append(
                {
                    "resume_id": str(doc["_id"]),
                    "uploaded_ts": uploaded_ts,
                    "label": label,
                }
            )
        logger.info("[resume_reader] listed %s resume(s) for user %s", len(items), user_oid)
        return items

    def get_resume_for_user(
        self,
        userid: Union[str, ObjectId],
        resume_id: Union[str, ObjectId],
    ) -> Dict[str, Any]:
        """Return one resume document for a user (includes parsed_data)."""
        logger.info("inside get_resume_for_user")
        user_oid = _as_user_object_id(userid)
        resume_oid = _as_user_object_id(resume_id)
        doc = self._collection().find_one(
            {"_id": resume_oid, "userid": user_oid},
            {"raw_text": 0},
        )
        if not doc:
            raise ValueError("resume not found for this user")

        ts = doc.get("uploaded_ts")
        parsed_data = doc.get("parsed_data")
        if not isinstance(parsed_data, dict):
            parsed_data = _default_parsed_data()

        resume_file = doc.get("resume_file") or {}
        filename = (resume_file.get("filename") or "").strip()
        name = (parsed_data.get("name") or "").strip()
        label = filename or name or "Resume"

        return {
            "resume_id": str(doc["_id"]),
            "userid": str(user_oid),
            "uploaded_ts": ts.isoformat() if hasattr(ts, "isoformat") else str(ts),
            "label": label,
            "parsed_data": parsed_data,
            "resume_file": resume_file,
        }

    def delete_all_resumes_for_user(self, userid: Union[str, ObjectId]) -> Tuple[int, int]:
        """
        Delete all resume documents and associated GridFS files for a user.

        Returns (resume_docs_deleted, gridfs_files_deleted).
        """
        logger.info("[resume_reader] delete_all_resumes_for_user")
        user_oid = _as_user_object_id(userid)
        self.ensure_resume_collection()
        cursor = self._collection().find(
            {"userid": user_oid},
            {"resume_file": 1},
        )
        gridfs_deleted = 0
        for doc in cursor:
            resume_file = doc.get("resume_file") or {}
            file_id_str = resume_file.get("file_id")
            if file_id_str:
                try:
                    self._fs.delete(ObjectId(file_id_str))
                    gridfs_deleted += 1
                except (InvalidId, Exception):
                    logger.warning(
                        "[resume_reader] could not delete GridFS file_id=%s",
                        file_id_str,
                    )
        result = self._collection().delete_many({"userid": user_oid})
        docs_deleted = int(result.deleted_count)
        logger.info(
            "[resume_reader] deleted %s resume(s) and %s GridFS file(s) for user %s",
            docs_deleted,
            gridfs_deleted,
            user_oid,
        )
        return docs_deleted, gridfs_deleted
